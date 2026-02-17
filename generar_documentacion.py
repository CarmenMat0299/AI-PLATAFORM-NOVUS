from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_heading_with_color(doc, text, level, color_rgb):
    """Agrega un encabezado con color personalizado"""
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_table_with_header(doc, headers, data):
    """Crea una tabla con encabezado formateado"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Datos
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

# Crear documento
doc = Document()

# Configurar márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ============================================================
# PORTADA
# ============================================================
title = doc.add_heading('JULIA', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.size = Pt(48)
run.font.color.rgb = RGBColor(0, 120, 212)  # Azure blue

subtitle = doc.add_paragraph('Plataforma de Inteligencia Artificial Conversacional')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(18)
subtitle.runs[0].font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
doc.add_paragraph()

# Logo/Empresa
company = doc.add_paragraph('Novus Soluciones S.A.')
company.alignment = WD_ALIGN_PARAGRAPH.CENTER
company.runs[0].font.size = Pt(16)
company.runs[0].font.bold = True

# Fecha
fecha = doc.add_paragraph(f'Fecha: {datetime.datetime.now().strftime("%d de %B de %Y")}')
fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
fecha.runs[0].font.size = Pt(12)

doc.add_page_break()

# ============================================================
# TABLA DE CONTENIDOS
# ============================================================
add_heading_with_color(doc, 'Tabla de Contenidos', 1, (0, 120, 212))
toc_items = [
    '1. Resumen Ejecutivo',
    '2. Arquitectura del Sistema',
    '3. Stack Tecnológico',
    '4. Componentes Principales',
    '5. Estructura del Proyecto',
    '6. Funcionalidades Clave',
    '7. Seguridad',
    '8. Métricas y Analytics',
    '9. Proceso de Desarrollo',
    '10. Conclusiones'
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. RESUMEN EJECUTIVO
# ============================================================
add_heading_with_color(doc, '1. Resumen Ejecutivo', 1, (0, 120, 212))

doc.add_paragraph(
    'JULIA es una plataforma de inteligencia artificial conversacional empresarial '
    'diseñada para automatizar la atención al cliente a través de múltiples canales '
    'de comunicación (WhatsApp Business y Microsoft Teams).'
)

add_heading_with_color(doc, '1.1 Objetivos del Proyecto', 2, (0, 100, 180))
objetivos = [
    'Automatizar la atención al cliente 24/7',
    'Reducir tiempos de respuesta a consultas frecuentes',
    'Escalar conversaciones complejas a agentes humanos',
    'Centralizar y monitorear todas las conversaciones',
    'Proporcionar métricas y analytics en tiempo real'
]
for obj in objetivos:
    doc.add_paragraph(obj, style='List Bullet')

add_heading_with_color(doc, '1.2 Alcance', 2, (0, 100, 180))
doc.add_paragraph(
    'La plataforma integra servicios de Azure Cognitive Services, '
    'incluyendo Azure OpenAI (GPT-4) para procesamiento de lenguaje natural, '
    'y proporciona un dashboard administrativo completo para monitoreo y gestión.'
)

doc.add_page_break()

# ============================================================
# 2. ARQUITECTURA DEL SISTEMA
# ============================================================
add_heading_with_color(doc, '2. Arquitectura del Sistema', 1, (0, 120, 212))

doc.add_paragraph(
    'JULIA utiliza una arquitectura moderna de tres capas: Frontend (React), '
    'Backend (FastAPI), y servicios en la nube (Azure).'
)

add_heading_with_color(doc, '2.1 Componentes Arquitectónicos', 2, (0, 100, 180))

componentes_arch = [
    'Canales de Entrada: WhatsApp Business y Microsoft Teams',
    'Backend API: FastAPI con Python 3.11',
    'Motor de IA: Azure OpenAI GPT-4',
    'Frontend: React 18 con Vite',
    'Deployment: Azure Container Apps con Docker',
    'Almacenamiento: Archivos JSON (conversations, escalations, users)',
    'Autenticación: JWT tokens con algoritmo HS256'
]
for comp in componentes_arch:
    doc.add_paragraph(comp, style='List Bullet')

add_heading_with_color(doc, '2.2 Flujo de Datos', 2, (0, 100, 180))
flujo = [
    'Cliente envía mensaje por WhatsApp o Teams',
    'Servicio de canal recibe y procesa el mensaje',
    'Se crea o actualiza la conversación en el sistema',
    'El mensaje se envía a Azure OpenAI (GPT-4)',
    'JULIA genera una respuesta contextual',
    'Si es necesario, JULIA usa herramientas (tools) para buscar información',
    'La respuesta se envía de vuelta al cliente',
    'Si JULIA no puede resolver, se crea una escalación',
    'Los administradores reciben notificaciones en el dashboard',
    'Las métricas se actualizan en tiempo real'
]
for i, paso in enumerate(flujo, 1):
    doc.add_paragraph(f'{i}. {paso}', style='List Number')

doc.add_page_break()

# ============================================================
# 3. STACK TECNOLÓGICO
# ============================================================
add_heading_with_color(doc, '3. Stack Tecnológico', 1, (0, 120, 212))

add_heading_with_color(doc, '3.1 Frontend', 2, (0, 100, 180))
frontend_data = [
    ['React', '18.3.1', 'Framework principal de UI'],
    ['Vite', '5.1.4', 'Build tool y dev server ultra rápido'],
    ['React Router DOM', '6.22.0', 'Navegación entre páginas'],
    ['Tailwind CSS', '3.4.1', 'Framework CSS utility-first'],
    ['Recharts', '2.15.4', 'Gráficas (Area, Bar, Pie)'],
    ['Lucide React', '0.344.0', 'Iconos modernos'],
    ['Axios', '1.6.7', 'Cliente HTTP']
]
add_table_with_header(doc, ['Tecnología', 'Versión', 'Propósito'], frontend_data)

doc.add_paragraph()

add_heading_with_color(doc, '3.2 Backend', 2, (0, 100, 180))
backend_data = [
    ['FastAPI', '0.104.1', 'Framework web moderno'],
    ['Uvicorn', '0.24.0', 'Servidor ASGI'],
    ['Python', '3.11', 'Lenguaje de programación'],
    ['Pydantic', '2.11.10', 'Validación de datos'],
    ['OpenAI SDK', '2.2.0', 'Cliente Azure OpenAI'],
    ['python-jose', '3.3.0', 'JWT tokens'],
    ['passlib + bcrypt', '1.7.4', 'Hashing de contraseñas'],
    ['python-dotenv', '1.0.0', 'Variables de entorno']
]
add_table_with_header(doc, ['Tecnología', 'Versión', 'Propósito'], backend_data)

doc.add_paragraph()

add_heading_with_color(doc, '3.3 Servicios de Azure', 2, (0, 100, 180))
azure_services = [
    'Azure OpenAI (GPT-4): Motor de inteligencia artificial',
    'Azure Speech Services: Text-to-Speech y Speech-to-Text',
    'Azure Vision: Análisis de imágenes',
    'Azure Cognitive Search: Búsqueda inteligente',
    'Azure Key Vault: Gestión de secretos',
    'Azure Communication Messages: WhatsApp Business API',
    'Azure Container Apps: Hosting serverless'
]
for service in azure_services:
    doc.add_paragraph(service, style='List Bullet')

doc.add_page_break()

# ============================================================
# 4. COMPONENTES PRINCIPALES
# ============================================================
add_heading_with_color(doc, '4. Componentes Principales del Sistema', 1, (0, 120, 212))

componentes = [
    {
        'titulo': '4.1 Motor de IA - JULIA',
        'descripcion': 'Utiliza Azure OpenAI GPT-4 para generar respuestas conversacionales naturales con contexto.',
        'funciones': [
            'Respuestas conversacionales naturales',
            'Context awareness (mantiene contexto)',
            'Integración con herramientas (tools/functions)',
            'Acceso a catálogo de servicios de Novus'
        ]
    },
    {
        'titulo': '4.2 Sistema de Conversaciones',
        'descripcion': 'Gestiona todas las conversaciones de múltiples canales.',
        'funciones': [
            'Gestión multi-canal (WhatsApp, Teams)',
            'Historial completo de mensajes',
            'Estados de conversación',
            'Métricas por conversación'
        ]
    },
    {
        'titulo': '4.3 Sistema de Escalaciones',
        'descripcion': 'Detecta cuando JULIA no puede resolver y escala a agentes humanos.',
        'funciones': [
            'Detección automática de limitaciones',
            'Creación de tickets',
            'Priorización de escalaciones',
            'Notificaciones a agentes'
        ]
    },
    {
        'titulo': '4.4 Autenticación y Usuarios',
        'descripcion': 'Sistema completo de gestión de usuarios y autenticación.',
        'funciones': [
            'Login/Logout con JWT',
            'Roles (Admin/Colaborador)',
            'Gestión de usuarios',
            'Reset de contraseñas por email'
        ]
    },
    {
        'titulo': '4.5 Notificaciones',
        'descripcion': 'Sistema de notificaciones en tiempo real.',
        'funciones': [
            'Notificaciones en dashboard',
            'Alertas de escalaciones',
            'Notificaciones por email'
        ]
    },
    {
        'titulo': '4.6 Métricas y Analytics',
        'descripcion': 'Monitoreo completo del rendimiento del sistema.',
        'funciones': [
            'Estadísticas en tiempo real',
            'Gráficas de actividad',
            'Distribución por canal',
            'Histórico de métricas'
        ]
    }
]

for comp in componentes:
    add_heading_with_color(doc, comp['titulo'], 2, (0, 100, 180))
    doc.add_paragraph(comp['descripcion'])
    doc.add_paragraph('Funcionalidades:')
    for func in comp['funciones']:
        doc.add_paragraph(func, style='List Bullet')
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 5. ESTRUCTURA DEL PROYECTO
# ============================================================
add_heading_with_color(doc, '5. Estructura del Proyecto', 1, (0, 120, 212))

doc.add_paragraph('El proyecto sigue una estructura organizada y modular:')

estructura = """
ai-platform-novus/
│
├── frontend/                    # Aplicación React
│   ├── src/
│   │   ├── components/         # Componentes reutilizables
│   │   ├── pages/              # Páginas principales
│   │   │   ├── Dashboard.jsx   # Panel de administración
│   │   │   ├── Chatview.jsx    # Vista de conversación
│   │   │   ├── Conversations.jsx
│   │   │   ├── Escalations.jsx
│   │   │   ├── TeamManagement.jsx
│   │   │   ├── Settings.jsx
│   │   │   └── Login.jsx
│   │   ├── context/            # Context API (Auth)
│   │   ├── services/           # API client (axios)
│   │   └── utils/              # Utilidades
│   ├── public/                 # Assets estáticos
│   └── package.json
│
├── src/                        # Backend Python
│   ├── api/
│   │   └── main.py            # FastAPI app principal
│   ├── services/              # Lógica de negocio (15+ servicios)
│   ├── models/                # Modelos Pydantic
│   ├── middleware/            # Middleware (CORS, Auth)
│   └── utils/                 # Utilidades
│
├── Dockerfile                 # Configuración de Docker
├── requirements.txt           # Dependencias Python
└── .env                       # Variables de entorno
"""

p = doc.add_paragraph(estructura)
p.style = 'No Spacing'
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# ============================================================
# 6. FUNCIONALIDADES CLAVE
# ============================================================
add_heading_with_color(doc, '6. Funcionalidades Clave', 1, (0, 120, 212))

add_heading_with_color(doc, '6.1 Para Clientes (WhatsApp/Teams)', 2, (0, 100, 180))
funciones_clientes = [
    'Conversación natural con JULIA 24/7',
    'Respuestas instantáneas a consultas',
    'Información sobre servicios de Novus',
    'Agendamiento de reuniones',
    'Búsqueda en catálogo de servicios'
]
for func in funciones_clientes:
    doc.add_paragraph(func, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '6.2 Para Administradores (Dashboard)', 2, (0, 100, 180))
funciones_admin = [
    'Vista consolidada de todas las conversaciones',
    'Métricas en tiempo real',
    'Gráficas de actividad (horaria, semanal, por canal)',
    'Sistema de escalaciones',
    'Gestión de usuarios y permisos',
    'Notificaciones de eventos importantes',
    'Monitoreo del estado del sistema',
    'Actividad reciente',
    'Dark mode y diseño responsive'
]
for func in funciones_admin:
    doc.add_paragraph(func, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '6.3 Páginas del Dashboard', 2, (0, 100, 180))
paginas_data = [
    ['Dashboard', '/', 'Métricas, gráficas, actividad, estado'],
    ['Conversaciones', '/conversations', 'Lista de conversaciones activas'],
    ['Chat', '/chat/:phone', 'Vista detallada de conversación'],
    ['Escalaciones', '/escalations', 'Tickets escalados a agentes'],
    ['Gestión de Equipo', '/team', 'CRUD de usuarios y roles'],
    ['Configuración', '/settings', 'Configuración del sistema'],
    ['Perfil', '/profile', 'Perfil del usuario'],
    ['Login', '/login', 'Autenticación']
]
add_table_with_header(doc, ['Página', 'Ruta', 'Descripción'], paginas_data)

doc.add_page_break()

# ============================================================
# 7. SEGURIDAD
# ============================================================
add_heading_with_color(doc, '7. Seguridad Implementada', 1, (0, 120, 212))

doc.add_paragraph('El sistema implementa múltiples capas de seguridad:')

seguridad = [
    'JWT Authentication: Tokens seguros con expiración configurable (480 minutos)',
    'Password Hashing: bcrypt para almacenamiento seguro de contraseñas',
    'CORS Protection: Configuración específica de dominios permitidos',
    'Role-based Access Control: Permisos por rol (Admin vs Colaborador)',
    'Environment Variables: Secretos no hardcodeados en el código',
    'Azure Key Vault: Gestión de secretos en producción',
    'HTTPS: Comunicación encriptada en producción',
    'Input Validation: Pydantic para validación de datos'
]
for item in seguridad:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 8. MÉTRICAS Y ANALYTICS
# ============================================================
add_heading_with_color(doc, '8. Métricas y Analytics', 1, (0, 120, 212))

doc.add_paragraph('El sistema rastrea las siguientes métricas en tiempo real:')

add_heading_with_color(doc, '8.1 Métricas Principales', 2, (0, 100, 180))
metricas = [
    'Conversaciones activas',
    'Mensajes procesados hoy',
    'Usuarios únicos',
    'Escalaciones pendientes',
    'Actividad por hora del día',
    'Distribución por canal (WhatsApp vs Teams)',
    'Histórico de 7 días',
    'Estado de servicios en tiempo real'
]
for metrica in metricas:
    doc.add_paragraph(metrica, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '8.2 Visualizaciones', 2, (0, 100, 180))
visualizaciones = [
    'Gráfica de Área: Actividad del día por hora',
    'Gráfica de Pastel: Distribución por canal',
    'Gráfica de Barras: Últimos 7 días por canal',
    'Tarjetas de estadísticas: Métricas principales',
    'Lista de actividad reciente: Últimas acciones'
]
for viz in visualizaciones:
    doc.add_paragraph(viz, style='List Bullet')

doc.add_page_break()

# ============================================================
# 9. PROCESO DE DESARROLLO
# ============================================================
add_heading_with_color(doc, '9. Proceso de Desarrollo', 1, (0, 120, 212))

doc.add_paragraph('El proyecto se desarrolló en cuatro fases principales:')

add_heading_with_color(doc, 'Fase 1: Backend API (Python + FastAPI)', 2, (0, 100, 180))
fase1 = [
    'Configuración de FastAPI y estructura de proyecto',
    'Integración con Azure OpenAI para el motor de JULIA',
    'Implementación de servicios (conversaciones, escalaciones, usuarios)',
    'Integración con WhatsApp Business API',
    'Integración con Microsoft Teams Bot Framework',
    'Sistema de autenticación JWT',
    'Endpoints de métricas y analytics'
]
for item in fase1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, 'Fase 2: Frontend (React + Vite)', 2, (0, 100, 180))
fase2 = [
    'Setup de Vite + React + Tailwind CSS',
    'Diseño del Dashboard con Recharts',
    'Vista de conversaciones estilo WhatsApp',
    'Sistema de autenticación (login, logout, roles)',
    'Gestión de escalaciones',
    'Gestión de usuarios (TeamManagement)',
    'Sistema de notificaciones',
    'Dark mode y responsive design'
]
for item in fase2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, 'Fase 3: Deployment (Docker + Azure)', 2, (0, 100, 180))
fase3 = [
    'Containerización con Docker',
    'Configuración de Azure Container Apps',
    'Variables de entorno y secretos',
    'Deploy a producción',
    'Configuración de dominio y SSL'
]
for item in fase3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, 'Fase 4: Refinamiento y Optimización', 2, (0, 100, 180))
fase4 = [
    'Optimización de gráficas y métricas',
    'Mejoras de UX/UI',
    'Auto-refresh y actualizaciones en tiempo real',
    'Correcciones de bugs (nombre "JULIA", mensajes largos, logo)',
    'Testing y validación'
]
for item in fase4:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 10. CONCLUSIONES
# ============================================================
add_heading_with_color(doc, '10. Conclusiones y Logros', 1, (0, 120, 212))

doc.add_paragraph(
    'JULIA representa una solución integral de inteligencia artificial conversacional '
    'que cumple con los estándares enterprise y está lista para producción.'
)

doc.add_paragraph()

add_heading_with_color(doc, '10.1 Logros Principales', 2, (0, 100, 180))
logros = [
    'Plataforma Multi-canal unificada (WhatsApp y Teams)',
    'IA Avanzada con GPT-4 y context awareness',
    'Dashboard Profesional con métricas en tiempo real',
    'Escalabilidad mediante Azure Container Apps (serverless)',
    'UI/UX Moderna con dark mode y diseño responsive',
    'Seguridad Enterprise con JWT, roles y Azure Key Vault',
    'Monitoreo Completo con métricas, logs y actividad'
]
for logro in logros:
    doc.add_paragraph(logro, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '10.2 Beneficios del Sistema', 2, (0, 100, 180))
beneficios = [
    'Reducción de tiempos de respuesta de horas a segundos',
    'Disponibilidad 24/7 sin intervención humana',
    'Escalabilidad automática según demanda',
    'Reducción de carga de trabajo en equipos de soporte',
    'Insights valiosos a través de métricas y analytics',
    'Experiencia de usuario moderna y profesional'
]
for beneficio in beneficios:
    doc.add_paragraph(beneficio, style='List Bullet')

doc.add_paragraph()

add_heading_with_color(doc, '10.3 Tecnologías Clave', 2, (0, 100, 180))
tech_table = [
    ['FastAPI', 'Framework backend moderno y rápido'],
    ['React + Vite', 'Desarrollo frontend con HMR instantáneo'],
    ['Azure OpenAI', 'GPT-4 enterprise con SLA'],
    ['Tailwind CSS', 'Diseño consistente y mantenible'],
    ['Docker + Azure', 'Deployment escalable y serverless']
]
add_table_with_header(doc, ['Tecnología', 'Beneficio'], tech_table)

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_para.add_run('© 2026 Novus Soluciones S.A. Todos los derechos reservados.')
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# Guardar documento
output_path = 'Documentacion_JULIA_AI_Platform.docx'
doc.save(output_path)
print(f'Documento creado exitosamente: {output_path}')
